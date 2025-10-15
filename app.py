from flask import Flask, render_template, request, jsonify, send_file
from flask import Flask, render_template, request, jsonify, session, redirect, url_for, make_response
import os
import cv2
import easyocr
import subprocess
from concurrent.futures import ThreadPoolExecutor, TimeoutError
from docx import Document
import numpy as np
from werkzeug.utils import secure_filename
from flask_cors import CORS
import shutil
import time
from threading import Lock

# ------------------ Flask app config ------------------
app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 500 * 1024 * 1024  # 500 MB
app.config['TEMPLATES_AUTO_RELOAD'] = True

# Enable CORS
CORS(app, resources={r"/*": {"origins": "*"}})

# Folders
UPLOAD_FOLDER = "uploads"
FRAME_FOLDER = "frames"
DOCX_FOLDER = "output"
DOCX_FILE = "captions.docx"
ALLOWED_EXTENSIONS = {".mp4", ".webm", ".avi", ".mov", ".mkv", ".flv"}
ALLOWED_MIMETYPES = {"video/mp4", "video/webm", "video/x-msvideo", "video/quicktime", 
                     "video/x-matroska", "video/x-flv"}

# Create folders
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(FRAME_FOLDER, exist_ok=True)
os.makedirs(DOCX_FOLDER, exist_ok=True)

# Global OCR reader and lock
ocr_reader = None
ocr_lock = Lock()

# ------------------ OCR setup ------------------
def initialize_ocr():
    """Initialize EasyOCR reader with error handling"""
    global ocr_reader
    if ocr_reader is None:
        try:
            print("Loading EasyOCR model... This may take a moment.")
            ocr_reader = easyocr.Reader(['en'], gpu=False, verbose=False)
            print("EasyOCR model loaded successfully!")
            return True
        except Exception as e:
            print(f"❌ Failed to load EasyOCR: {e}")
            return False
    return True

# ------------------ Routes ------------------

@app.route("/", methods=["GET"])
def index():
    """Main video upload page"""
    return render_template("index.html")

@app.route("/login", methods=["GET"])
def login():
    """Login page"""
    return render_template("login.html")

@app.route("/signup", methods=["GET"])
def signup():
    """Signup page"""
    return render_template("signup.html")

@app.route("/upload", methods=["POST"])
def upload():
    """Handle video upload and processing"""
    try:
        print("\n" + "="*50)
        print("NEW UPLOAD REQUEST")
        print("="*50)
        
        # Check OCR is initialized
        if not initialize_ocr():
            return jsonify({
                "status": "error", 
                "error": "OCR system not available. Please try again later."
            }), 500
        
        # Cleanup old files
        print("Cleaning up old files...")
        clear_folder(FRAME_FOLDER)
        clear_folder(UPLOAD_FOLDER)
        docx_path = os.path.join(DOCX_FOLDER, DOCX_FILE)
        safe_remove_file(docx_path)

        # Validate request
        if 'video' not in request.files:
            print("❌ Error: No video file in request")
            return jsonify({"status": "error", "error": "No video file provided"}), 400

        video = request.files["video"]
        if not video.filename:
            print("❌ Error: Empty filename")
            return jsonify({"status": "error", "error": "No file selected"}), 400
        
        # Validate MIME type
        if video.mimetype not in ALLOWED_MIMETYPES:
            print(f"❌ Error: Invalid MIME type: {video.mimetype}")
            return jsonify({
                "status": "error", 
                "error": f"Unsupported video format. Please upload MP4, WebM, AVI, MOV, or MKV files."
            }), 400
            
        if not allowed_file(video.filename):
            print(f"❌ Error: Unsupported format: {video.filename}")
            return jsonify({"status": "error", "error": "Unsupported file format"}), 400

        # Save uploaded file
        raw_path = os.path.join(UPLOAD_FOLDER, secure_filename(video.filename))
        video_path = os.path.join(UPLOAD_FOLDER, "input.mp4")
        
        print(f"📥 Saving uploaded file: {video.filename}")
        video.save(raw_path)
        
        # Validate file size
        if not os.path.exists(raw_path) or os.path.getsize(raw_path) == 0:
            return jsonify({"status": "error", "error": "Failed to save video file"}), 500
            
        file_size_mb = os.path.getsize(raw_path) / (1024 * 1024)
        print(f"   File size: {file_size_mb:.2f} MB")

        # Check FFmpeg availability
        if not check_ffmpeg():
            return jsonify({
                "status": "error",
                "error": "FFmpeg is not installed. Please install FFmpeg to process videos."
            }), 500

        # Compress video with FFmpeg
        print("🎬 Compressing video with FFmpeg...")
        try:
            result = subprocess.run([
                "ffmpeg", "-i", raw_path,
                "-vf", "scale='if(gt(iw,1920),1920,iw)':'if(gt(ih,1080),1080,ih)'",
                "-c:v", "libx264", "-preset", "fast", "-crf", "22",
                "-y", video_path
            ], stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=300, check=False)

            if result.returncode != 0:
                error_msg = result.stderr.decode('utf-8', errors='ignore')[-500:]
                print(f"❌ FFmpeg failed: {error_msg}")
                return jsonify({
                    "status": "error",
                    "error": "Video compression failed. Please try a different video format."
                }), 500
        except subprocess.TimeoutExpired:
            return jsonify({
                "status": "error",
                "error": "Video processing timeout. Please try a shorter video."
            }), 500

        if not os.path.exists(video_path) or os.path.getsize(video_path) == 0:
            return jsonify({
                "status": "error",
                "error": "Video compression produced invalid output."
            }), 500

        compressed_size_mb = os.path.getsize(video_path) / (1024 * 1024)
        print(f"✅ Compression complete: {compressed_size_mb:.2f} MB")

        # Extract frames
        print("🖼️  Extracting frames...")
        try:
            frame_count = extract_frames(video_path)
            print(f"✅ Extracted {frame_count} frames")
        except Exception as e:
            print(f"❌ Frame extraction failed: {e}")
            return jsonify({
                "status": "error",
                "error": f"Frame extraction failed: {str(e)}"
            }), 500

        # Run OCR
        print("🔍 Running OCR on frames...")
        try:
            caption_count = run_ocr_parallel()
            print(f"✅ Extracted {caption_count} captions")
        except Exception as e:
            print(f"❌ OCR processing failed: {e}")
            return jsonify({
                "status": "error",
                "error": f"Caption extraction failed: {str(e)}"
            }), 500

        # Cleanup temporary files
        safe_remove_file(raw_path)
        safe_remove_file(video_path)

        # Success!
        print("="*50)
        print("✅ UPLOAD COMPLETED SUCCESSFULLY")
        print("="*50 + "\n")
        
        return jsonify({
            "status": "success",
            "message": f"Successfully extracted {caption_count} captions",
            "download_url": "/download"
        }), 200

    except Exception as e:
        print(f"❌ UNEXPECTED ERROR: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "status": "error",
            "error": f"Processing failed: {str(e)}"
        }), 500

@app.route("/download", methods=["GET"])
def download():
    """Download the generated captions file"""
    docx_path = os.path.join(DOCX_FOLDER, DOCX_FILE)
    if not os.path.exists(docx_path):
        return jsonify({"status": "error", "error": "No captions file available"}), 404
    
    return send_file(
        docx_path,
        as_attachment=True,
        download_name="captions.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# ------------------ Error Handlers ------------------

@app.errorhandler(413)
def too_large(e):
    return jsonify({"status": "error", "error": "File too large. Max 500MB."}), 413

@app.errorhandler(404)
def not_found(e):
    if request.path in ['/upload', '/download'] or request.path.startswith('/api'):
        return jsonify({"status": "error", "error": "Endpoint not found"}), 404
    try:
        return render_template("index.html")
    except:
        return jsonify({"status": "error", "error": "Page not found"}), 404

@app.errorhandler(500)
def internal_error(e):
    print(f"Internal server error: {str(e)}")
    return jsonify({"status": "error", "error": "Internal server error"}), 500

# ------------------ Helper Functions ------------------

def check_ffmpeg():
    """Check if FFmpeg is available"""
    try:
        result = subprocess.run(
            ["ffmpeg", "-version"],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=5,
            check=False
        )
        return result.returncode == 0
    except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
        return False

def allowed_file(filename):
    """Check if file extension is allowed"""
    return os.path.splitext(filename)[1].lower() in ALLOWED_EXTENSIONS

def safe_remove_file(filepath):
    """Safely remove a file with retries"""
    if not os.path.exists(filepath):
        return
    for attempt in range(3):
        try:
            os.remove(filepath)
            return
        except Exception as e:
            if attempt == 2:
                print(f"Warning: Could not delete {filepath}: {e}")
            time.sleep(0.1)

def clear_folder(folder):
    """Delete all files in a folder with better error handling"""
    if not os.path.exists(folder):
        return
    for f in os.listdir(folder):
        file_path = os.path.join(folder, f)
        if os.path.isfile(file_path):
            safe_remove_file(file_path)

def preprocess_image(img):
    """Preprocess image for better OCR results"""
    if img is None or img.size == 0:
        raise ValueError("Invalid image for preprocessing")
    
    # Convert to grayscale
    if len(img.shape) == 3:
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    else:
        gray = img
    
    # Apply Gaussian blur
    blur = cv2.GaussianBlur(gray, (3, 3), 0)
    
    # Apply adaptive threshold
    thresh = cv2.adaptiveThreshold(
        blur, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
    )
    
    return thresh

def detect_dynamic_caption_region(frames):
    """Detect the region where captions typically appear"""
    bboxes = []
    
    for frame in frames[:min(5, len(frames))]:  # Limit to first 5 frames
        try:
            pre = preprocess_image(frame)
            with ocr_lock:
                result = ocr_reader.readtext(pre)
            
            for (bbox, text, conf) in result:
                if conf > 0.5 and len(text.strip()) > 0:
                    x_coords = [pt[0] for pt in bbox]
                    y_coords = [pt[1] for pt in bbox]
                    x1, x2 = int(min(x_coords)), int(max(x_coords))
                    y1, y2 = int(min(y_coords)), int(max(y_coords))
                    bboxes.append((x1, y1, x2, y2))
        except Exception as e:
            print(f"   Warning: Region detection failed on frame: {e}")
            continue
    
    if not bboxes:
        return None
    
    # Return bounding box that encompasses all detected text
    return (
        max(0, min(b[0] for b in bboxes) - 10),
        max(0, min(b[1] for b in bboxes) - 10),
        max(b[2] for b in bboxes) + 10,
        max(b[3] for b in bboxes) + 10
    )

def extract_frames(video_path):
    """Extract frames from video at regular intervals"""
    cap = cv2.VideoCapture(video_path)
    
    if not cap.isOpened():
        raise Exception("Failed to open video file. The file may be corrupted.")
    
    fps = cap.get(cv2.CAP_PROP_FPS)
    total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
    
    if fps == 0 or total_frames == 0:
        cap.release()
        raise Exception("Invalid video file. Cannot read video properties.")
    
    duration = min(int(total_frames / fps), 3600)  # Max 1 hour
    
    # Sample first few frames to detect caption region
    sample_frames = []
    sec = 0
    while sec < min(duration, 20) and len(sample_frames) < 5:
        cap.set(cv2.CAP_PROP_POS_MSEC, sec * 1000)
        ret, frame = cap.read()
        if ret and frame is not None and frame.size > 0:
            sample_frames.append(frame)
        sec += 5
    
    region = None
    if sample_frames:
        region = detect_dynamic_caption_region(sample_frames)
        if region:
            print(f"   Caption region detected: {region}")
        else:
            print("   No specific caption region detected, using full frame")
    
    # Reset to beginning
    cap.set(cv2.CAP_PROP_POS_MSEC, 0)
    
    # Extract frames for OCR
    sec = 0
    frame_interval = max(2, min(10, duration // 120))
    frame_count = 0
    
    while sec < duration:
        cap.set(cv2.CAP_PROP_POS_MSEC, sec * 1000)
        ret, frame = cap.read()
        
        if not ret or frame is None or frame.size == 0:
            sec += frame_interval
            continue
        
        try:
            # Crop to caption region if detected
            if region:
                h, w = frame.shape[:2]
                x1, y1, x2, y2 = region
                # Validate crop coordinates
                x1, y1 = max(0, x1), max(0, y1)
                x2, y2 = min(w, x2), min(h, y2)
                if x2 > x1 and y2 > y1:
                    cropped = frame[y1:y2, x1:x2]
                else:
                    cropped = frame
            else:
                cropped = frame
            
            filename = os.path.join(FRAME_FOLDER, f"frame_{sec:04d}.png")
            success = cv2.imwrite(filename, cropped)
            if success:
                frame_count += 1
        except Exception as e:
            print(f"   Warning: Failed to save frame at {sec}s: {e}")
        
        sec += frame_interval
    
    cap.release()
    
    if frame_count == 0:
        raise Exception("No frames could be extracted from the video")
    
    return frame_count

def run_ocr_parallel():
    """Run OCR on all extracted frames in parallel"""
    doc = Document()
    doc.add_heading("Extracted Captions", level=1)
    
    frame_files = sorted(os.listdir(FRAME_FOLDER))
    if not frame_files:
        raise Exception("No frames found to process")
    
    last_caption = ""
    last_time = -5
    caption_count = 0
    results = []
    
    # Process frames with timeout
    with ThreadPoolExecutor(max_workers=4) as executor:
        try:
            futures = list(executor.map(process_frame, frame_files, timeout=300))
            results = [r for r in futures if r is not None]
        except TimeoutError:
            print("   ⚠️  OCR processing timeout, using partial results")
            results = []
    
    # Sort results by timestamp
    results.sort(key=lambda x: x[0] if x else 0)
    
    for item in results:
        if item:
            time_sec, text = item
            # Only add if different from last caption and not empty
            if text and text != last_caption:
                prefix = f"[{time_sec:02d}s] " if time_sec - last_time >= 3 else ""
                doc.add_paragraph(f"{prefix}{text}")
                last_caption = text
                last_time = time_sec
                caption_count += 1
    
    if caption_count == 0:
        doc.add_paragraph("No captions detected in video.")
    
    docx_path = os.path.join(DOCX_FOLDER, DOCX_FILE)
    doc.save(docx_path)
    
    return caption_count

def process_frame(fname):
    """Process a single frame with OCR"""
    try:
        img_path = os.path.join(FRAME_FOLDER, fname)
        
        if not os.path.exists(img_path):
            return None
            
        img = cv2.imread(img_path)
        
        if img is None or img.size == 0:
            print(f"   Warning: Could not read {fname}")
            return None
        
        pre = preprocess_image(img)
        
        # Use lock to prevent concurrent OCR calls
        with ocr_lock:
            result = ocr_reader.readtext(pre)
        
        # Extract text with confidence > 0.6
        texts = [t.strip() for _, t, conf in result if conf > 0.6 and t.strip()]
        text = " ".join(texts)
        
        if text:
            # Extract timestamp from filename
            time_sec = int(fname.split('_')[1].split('.')[0])
            return (time_sec, text)
            
    except Exception as e:
        print(f"   ❌ OCR failed on {fname}: {e}")
    
    return None


if __name__ == "__main__":
    print("\n" + "="*60)
    print("🎬 EduCaption - Video Caption Extractor")
    print("="*60)
    print("\n📋 Checklist:")
    print("   ✅ Flask app configured")
    
    # Check templates folder
    if os.path.exists('templates') and os.path.exists('templates/index.html'):
        print("   ✅ Templates folder found")
    else:
        print("   ❌ Templates folder missing!")
        print("      Create 'templates/' folder and add 'index.html'")
    
    # Check FFmpeg
    if check_ffmpeg():
        print("   ✅ FFmpeg is installed")
    else:
        print("   ⚠️  FFmpeg not found - video processing will fail")
    
    # Initialize OCR
    if initialize_ocr():
        print("   ✅ EasyOCR initialized")
    else:
        print("   ⚠️  EasyOCR initialization failed")
    
    print("\n🚀 Starting server on http://localhost:5000")
    print("="*60 + "\n")
    
    app.run(debug=True, host="localhost", port=5000, threaded=True)
